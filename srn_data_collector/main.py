import argparse
import json
import os
import re
from collections import defaultdict
from fnmatch import fnmatch
from typing import Dict, List

import pandas as pd
import requests
import yaml
from docx import Document
from easynmt import EasyNMT
from rapidfuzz import fuzz, process
from sqlalchemy import MetaData, Table, create_engine, select
from sqlalchemy.orm import Session
from stqdm import stqdm
from tqdm import tqdm

from srn_data_collector.annotations_utils.collect_api_data import (
    get_compliance_item_instance,
    get_financial_index,
    get_reporting_requirement_instance,
    get_srn_companies,
    get_srn_documents,
    get_standard_list,
    get_standard_requirements,
    get_values_with_revisions,
)
from srn_data_collector.annotations_utils.data_model import (
    Base,
    BlobLvlAnnotations,
    BlobLvlAnnotationsModel,
    CompanyDetailsModel,
    CompanyDetailsTable,
    ComplianceItems,
    ComplianceItemsModel,
    DocumentInstances,
    DocumentInstancesModel,
    EsrsReqMapping,
    EsrsReqMappingModel,
    Indices,
    IndicesModel,
    ReportingRequirements,
    ReportingRequirementsModel,
    RptRequirementsMapping,
    RptRequirementsMappingModel,
    StandardsList,
    StandardsListModel,
    ValuesWithRevisions,
    ValuesWithRevisionsModel,
)
from srn_data_collector.annotations_utils.llm_utils_helper import (
    estimate_api_cost,
    json_translate_estimate_gpt,
    predict_samples_parallel,
    translate_json_parallel,
)
from srn_data_collector.annotations_utils.logger_setup import setup_logger
from srn_data_collector.viper_parser import get_page_dict, process_annotation_idx


def collect_companies_data(session):
    companies = get_srn_companies()
    for company in tqdm(companies, desc="Processing companies"):
        company_fin_idx_ids = company.pop("indices")
        indices = []
        for company_fin_idx_id in company_fin_idx_ids:
            fin_idx_data = get_financial_index(company_fin_idx_id)
            indices_model = IndicesModel(fin_idx_id=None, **fin_idx_data)
            indices.append(
                Indices(
                    fin_idx_id=indices_model.id,
                    name=indices_model.name,
                    name_short=indices_model.name_short,
                    country=indices_model.country,
                    coverage=indices_model.coverage,
                    updated=indices_model.updated,
                    href=indices_model.href,
                )
            )

        company["indices"] = indices
        company_details_model = CompanyDetailsModel(**company)
        company_details_table = CompanyDetailsTable(
            id=company_details_model.id,
            name=company_details_model.name,
            isin=company_details_model.isin,
            country=company_details_model.country,
            sector=company_details_model.sector,
            href=company_details_model.href,
            href_logo=company_details_model.href_logo,
            company_type=company_details_model.company_type,
            indices=indices,
        )
        session.add(company_details_table)

    # Commit the changes to the database
    session.commit()


def get_custom_revisions(revision_orig):
    revisions_modified = []
    verified_revision = revision_orig.pop("verified_revision")
    all_revisions = revision_orig.pop("revisions")
    verified_at = revision_orig.pop("verified_at")
    for revision in tqdm(all_revisions, desc="Collecting custom revisions"):
        revision_obj = {}
        if verified_revision and (revision["id"] == verified_revision["id"]):
            # For all verified revisions
            revision_obj.update({"is_verified": True})
            revision_obj.update({"verified_at": verified_at})
        else:
            # For all unverified revisions
            revision_obj.update({"is_verified": False})
            revision_obj.update({"verified_at": None})
        revision_obj.update(revision_orig)
        revision_obj.update(revision)

        revisions_modified.append(revision_obj)

    return revisions_modified


def collect_annotations(session):
    companies = session.query(CompanyDetailsTable).all()
    for company in tqdm(companies, desc="Retrieving companies annotations"):
        revisions_all: List[Dict] = get_values_with_revisions(company.id)
        for revision in revisions_all:
            modified_revisions = get_custom_revisions(revision)
            for modified_revision in modified_revisions:
                values_with_revision_model = ValuesWithRevisionsModel(**modified_revision)
                if not session.query(ValuesWithRevisions).filter_by(id=values_with_revision_model.id).all():
                    values_with_revisions_table = ValuesWithRevisions(
                        id=values_with_revision_model.id,
                        company_id=values_with_revision_model.company_id,
                        compliance_item_id=values_with_revision_model.compliance_item_id,
                        year=values_with_revision_model.year,
                        verified_at=values_with_revision_model.verified_at,
                        company_value_id=values_with_revision_model.company_value_id,
                        created_at=values_with_revision_model.created_at,
                        review_comment=values_with_revision_model.review_comment,
                        flagged_for_review=values_with_revision_model.flagged_for_review,
                        exists=values_with_revision_model.exists,
                        value=values_with_revision_model.value,
                        own_calculation=values_with_revision_model.own_calculation,
                        document_id=values_with_revision_model.document_id,
                        document_ref=values_with_revision_model.document_ref,
                        material=values_with_revision_model.material,
                        material_reason=values_with_revision_model.material_reason,
                        is_verified=values_with_revision_model.is_verified,
                    )
                    session.add(values_with_revisions_table)

                    # Commit the changes to the database
                    session.commit()


def get_storage_path(doc_data_model, base_path):
    file_path = os.path.join(base_path, doc_data_model.company_id, doc_data_model.year, doc_data_model.type)
    file_name_abs = os.path.join(file_path, f"{doc_data_model.id}{os.path.splitext(doc_data_model.href)[-1]}")
    # if not os.path.exists(file_path):
    #     os.makedirs(file_path)

    if os.path.isfile(file_name_abs):
        # If the file exists that means the pdf is downloaded already
        # download_document(file_name_abs, doc_data_model.href)
        return file_name_abs


def get_all_files_in_directory(directory, wildcard=None):
    all_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            if wildcard and not fnmatch(file, wildcard):
                continue  # Skip files that don't match the wildcard
            all_files.append(file_path)
    return all_files


def collect_documents_metadata(session, local_storage_path):
    for file_path in tqdm(
        get_all_files_in_directory(local_storage_path, wildcard="*.pdf"), desc="Processing pdf files"
    ):
        filename_without_extension = os.path.basename(file_path).split(".")
        if filename_without_extension:
            document_id = filename_without_extension[0]
            document_data = get_srn_documents(document_id)

            if "detail" not in document_data:
                document_instances_model = DocumentInstancesModel(**document_data)

                # Check if the document exists already
                if not session.query(DocumentInstances).filter_by(id=document_instances_model.id).all():
                    wz_storage_location = get_storage_path(document_instances_model, base_path=local_storage_path)
                    if wz_storage_location:
                        document_instances_table = DocumentInstances(
                            id=document_instances_model.id,
                            name=document_instances_model.name,
                            href=document_instances_model.href,
                            type=document_instances_model.type,
                            year=document_instances_model.year,
                            company_id=document_instances_model.company_id,
                            created_at=document_instances_model.created_at,
                            created_by_info=document_instances_model.created_by_info,
                            wz_storage_location=wz_storage_location,
                        )
                        session.add(document_instances_table)
                        # Commit the changes to the database
                        session.commit()


def collect_compliance_item_data(session):
    compliance_items = get_compliance_item_instance()
    for compliance_item in tqdm(compliance_items, desc="Retrieving compliance item data for each annotation"):
        compliance_item_model = ComplianceItemsModel(**compliance_item)
        if not session.query(ComplianceItems).filter_by(id=compliance_item_model.id).all():
            compliance_items_table = ComplianceItems(
                id=compliance_item_model.id,
                reporting_requirement_id=compliance_item_model.reporting_requirement_id,
                name=compliance_item_model.name,
                status=compliance_item_model.status,
            )
            session.add(compliance_items_table)

            # Commit the changes to the database
            session.commit()


def collect_reporting_item_data(session):
    compliance_items = session.query(ComplianceItems).all()
    for compliance_item in tqdm(compliance_items, desc="Retrieving reporting items for corresponding compliance item"):
        reporting_req_data = get_reporting_requirement_instance(compliance_item.reporting_requirement_id)
        rpt_requirement_model = ReportingRequirementsModel(**reporting_req_data)
        if not session.query(ReportingRequirements).filter_by(id=rpt_requirement_model.id).all():
            rpt_req_table = ReportingRequirements(
                id=rpt_requirement_model.id,
                name=rpt_requirement_model.name,
                topic1=rpt_requirement_model.topic1,
                topic2=rpt_requirement_model.topic2,
                topic3=rpt_requirement_model.topic3,
                topic4=rpt_requirement_model.topic4,
                req_unit=rpt_requirement_model.req_unit,
                req_time=rpt_requirement_model.req_time,
                req_descriptor=rpt_requirement_model.req_descriptor,
                sorting_id=rpt_requirement_model.sorting_id,
            )
            session.add(rpt_req_table)

            # Commit the changes to the database
            session.commit()


def get_page_data(page_number: int, json_fpath: str, logger, error_stat_dict=None) -> List[Dict]:
    page_data_dict = None
    try:
        with open(json_fpath, "r") as f:
            all_data_dict = json.load(f)
        page_data_dict = all_data_dict[str(page_number)]
    except KeyError:
        if page_number > len(all_data_dict) or page_number == 0:
            logger.error("\033[91mPage is not available in the pdf file\033[0m")
            if "PageNotFoundError" not in error_stat_dict:
                error_stat_dict["PageNotFoundError"] = 0
            error_stat_dict["PageNotFoundError"] += 1
    except FileNotFoundError:
        if "FileNotFoundError" not in error_stat_dict:
            error_stat_dict["FileNotFoundError"] = []
        filename = os.path.splitext(os.path.basename(json_fpath))[0]
        error_stat_dict["FileNotFoundError"].append(filename)
        logger.error(f"\033[91mDocument {filename} is in the database but not available\033[0m")
    return page_data_dict, error_stat_dict


def fuzzy_match_strings(line: str, tgt_list: List[str], threshold=97):
    return process.extractOne(line, tgt_list, scorer=fuzz.QRatio, score_cutoff=threshold)


def gpt_match_strings(samples_list, reponse_storage_path):
    with stqdm(desc="**Processing requirement**", total=len(samples_list), backend=True, frontend=True) as pbar:
        results_list = predict_samples_parallel(samples_list, pbar, storage_path=reponse_storage_path)
    return results_list


def collect_blob_lvl_annotations(
    session,
    logger,
    local_storage_path,
    fuzzy_match_thresh,
    error_stat_dict=None,
    estimate_cost=False,
    gpt_reponse_file_path=None,
):
    blob_type_ignore_list = ["header/footer", "headline"]
    cost_est_dict = {"api_len": 0, "api_cost": 0}
    sample_list = []

    get_annotations_query = (
        session.query(ValuesWithRevisions)
        .filter(ValuesWithRevisions.value.notin_(["", "\\", "/"]))
        .filter(
            ValuesWithRevisions.document_ref.notin_(
                [
                    "",
                    "\\",
                    "/",
                    "from the previous data",
                    "N.a.",
                    "n.p.",
                    "n/a",
                    "data_too_long_to_display",
                    "N.a.",
                    "na",
                    "n.a.",
                    "NA",
                ]
            )
        )
    )
    annotations = get_annotations_query.all()
    existing_annotations = [i.revision_id for i in session.query(BlobLvlAnnotations).all()]
    for annotation in tqdm(annotations, desc="Extracting annotations from json files"):
        if annotation.id not in existing_annotations:
            doc_id = annotation.document_id
            document_ref, error_stat_dict = process_annotation_idx(
                annotation.document_ref, error_stat_dict=error_stat_dict
            )

            if document_ref and annotation.value:
                company_id = annotation.company_id

                # get document metadata
                docs_meta_data = session.query(DocumentInstances).filter_by(id=doc_id).all()
                if not docs_meta_data:
                    logger.error(f"Data for {doc_id} not available!!")
                    continue
                docs_meta_data = docs_meta_data[0]

                # get page parsed json data
                parsed_data_file_path = os.path.join(
                    local_storage_path, company_id, docs_meta_data.year, docs_meta_data.type, f"{doc_id}.json"
                )

                # go page by page in the document
                for doc_page_num in document_ref:
                    blobs, error_stat_dict = get_page_data(doc_page_num, parsed_data_file_path, logger, error_stat_dict)
                    if blobs:
                        # for fuzzy matching
                        if fuzzy_match_thresh:
                            blobs_value_list = [
                                blob["text"].lower()
                                for blob in blobs
                                if blob["class_name"] not in blob_type_ignore_list
                            ]
                            matched_blob = fuzzy_match_strings(
                                annotation.value.lower(), blobs_value_list, threshold=fuzzy_match_thresh
                            )
                            if matched_blob:
                                update_blob_data(**matched_blob, session=session, error_stat_dict=error_stat_dict)

                        # for gpt api matching
                        else:
                            compliance_item_id = (
                                session.query(ComplianceItems).filter_by(id=annotation.compliance_item_id).all()
                            )
                            if estimate_cost:
                                api_len, api_cost = estimate_api_cost(
                                    compliance_item_id[0].name, annotation.value.lower(), blobs, max_tokens=256
                                )
                                cost_est_dict["api_len"] += api_len
                                cost_est_dict["api_cost"] += api_cost
                            else:
                                # create the sample list for gpt api matchin
                                sample_list.append(
                                    {
                                        "compliance_item": compliance_item_id[0].name,
                                        "annotation": annotation,
                                        "blobs": blobs,
                                    }
                                )
    # process the sample list for gpt api matching
    if not estimate_cost and not fuzzy_match_thresh:
        matched_blob_list: List[Dict] = gpt_match_strings(sample_list, gpt_reponse_file_path)

        for matched_blob in tqdm(matched_blob_list, desc="updating database with matched blobs"):
            if matched_blob:
                update_blob_data(**matched_blob, session=session, error_stat_dict=error_stat_dict)

    return error_stat_dict, cost_est_dict


def update_blob_data(annotation, blobs, matched_blob, session, error_stat_dict):
    if matched_blob:
        for blob_idx in matched_blob:
            try:
                blob_idx = int(blob_idx)
                if blob_idx < len(blobs):
                    blob_data = blobs[blob_idx]
                    annotated_data_dict = {
                        "blob_id": blob_idx,
                        "revision_id": annotation.id,
                        "document_id": annotation.document_id,
                        "blob_start_id": None,
                        "blob_end_id": None,
                        "blob_class_id": blob_data["class_id"],
                        "blob_class_name": blob_data["class_name"],
                        "blob_text": blob_data["text"],
                        "blob_box_x1": blob_data["box"][0],
                        "blob_box_y1": blob_data["box"][1],
                        "blob_box_x2": blob_data["box"][2],
                        "blob_box_y2": blob_data["box"][3],
                        "annotation_status": True,
                    }
                    blob_lvl_annotation_model = BlobLvlAnnotationsModel(**annotated_data_dict)
                    blob_lvl_annotations_table = BlobLvlAnnotations(
                        blob_id=blob_lvl_annotation_model.blob_id,
                        revision_id=blob_lvl_annotation_model.revision_id,
                        document_id=blob_lvl_annotation_model.document_id,
                        blob_start_id=blob_lvl_annotation_model.blob_start_id,
                        blob_end_id=blob_lvl_annotation_model.blob_end_id,
                        blob_class_id=blob_lvl_annotation_model.blob_class_id,
                        blob_class_name=blob_lvl_annotation_model.blob_class_name,
                        blob_text=blob_lvl_annotation_model.blob_text,
                        blob_box_x1=blob_lvl_annotation_model.blob_box_x1,
                        blob_box_y1=blob_lvl_annotation_model.blob_box_y1,
                        blob_box_x2=blob_lvl_annotation_model.blob_box_x2,
                        blob_box_y2=blob_lvl_annotation_model.blob_box_y2,
                        annotation_status=blob_lvl_annotation_model.annotation_status,
                    )
                    session.add(blob_lvl_annotations_table)
                    # Commit the changes to the database
                    session.commit()
                else:
                    print(f"{blob_idx} index is out of range")
            except TypeError:
                if "DBErrorTypeError" not in error_stat_dict:
                    error_stat_dict["DBErrorTypeError"] = 0
                error_stat_dict["DBErrorTypeError"] += 1
            except ValueError:
                if "DBErrorValueError" not in error_stat_dict:
                    error_stat_dict["DBErrorValueError"] = 0
                error_stat_dict["DBErrorValueError"] += 1
        return error_stat_dict


def collect_standards_list(session):
    std_list = get_standard_list()
    for std in tqdm(std_list, desc="Updating standards list table"):
        standards_list = StandardsListModel(**std)
        if not session.query(StandardsList).filter_by(id=standards_list.id).all():
            std_list_table = StandardsList(
                id=standards_list.id,
                name=standards_list.name,
                version=standards_list.version,
                family=standards_list.family,
                href=standards_list.href,
                topic=standards_list.topic,
            )
            session.add(std_list_table)

            # Commit the changes to the database
            session.commit()


def collect_std_requirements(session):
    std_requirements = get_standard_requirements()
    for std_requirement in tqdm(std_requirements, desc="Creating RptRequirementsMapping table"):
        if "standard" in std_requirement:
            std_requirement["standard"] = std_requirement["standard"]["id"]
            reporting_req_map_model = RptRequirementsMappingModel(**std_requirement)
            if not session.query(RptRequirementsMapping).filter_by(id=reporting_req_map_model.id).all():
                reporting_req_map_table = RptRequirementsMapping(
                    id=reporting_req_map_model.id,
                    reporting_requirement_id=reporting_req_map_model.reporting_requirement_id,
                    standard=reporting_req_map_model.standard,
                    exists=reporting_req_map_model.exists,
                    source=reporting_req_map_model.source,
                    comment=reporting_req_map_model.comment,
                    type=reporting_req_map_model.type,
                )
                session.add(reporting_req_map_table)
                # Commit the changes to the database
                session.commit()


def run_blob_conversion(session, model_name, local_storage_path, estimate_cost: bool = False, limit_docs: int = 10):
    """
    This function can estimate the cost for translation task and run the translation
    """
    estimates_json = {"api_len": 0, "api_cost": 0}
    json_files_list = get_all_files_in_directory(local_storage_path, wildcard="*.json")
    # existing_docs_list = [doc.id for doc in session.query(DocumentInstances)]
    annotated_docs_list = set([doc.document_id for doc in session.query(BlobLvlAnnotations)])
    json_files_list = [
        json_file
        for json_file in json_files_list
        if os.path.splitext(os.path.basename(json_file))[0] in annotated_docs_list
    ]

    if limit_docs:
        json_files_list = json_files_list[:limit_docs]
    # json_translate_estimate
    for json_file in tqdm(json_files_list, desc="Processing jsons for translation task"):
        if not os.path.exists(f"{os.path.splitext(json_file)[0]}_de.json"):
            with open(json_file, "r") as f:
                doc_dict = json.load(f)

            blob_text_list = []
            for _, page_blobs in tqdm(
                doc_dict.items(), desc=f"Translating blobs for : {os.path.splitext(os.path.basename(json_file))[0]}"
            ):
                for blob in page_blobs:
                    if estimate_cost:
                        api_len, api_cost = json_translate_estimate_gpt(blob["text"], max_tokens=2048)
                        estimates_json["api_len"] += api_len
                        estimates_json["api_cost"] += api_cost
                    else:
                        blob_text_list.append(blob)

            if not estimate_cost:
                with stqdm(desc="**Processing blobs**", total=len(blob_text_list), backend=True, frontend=True) as pbar:
                    translate_json_parallel(model_name, blob_text_list, pbar=pbar, max_tokens=2048)

            with open(f"{os.path.splitext(json_file)[0]}_de.json", "w") as f:
                json.dump(doc_dict, f, indent=2)
    return estimates_json


def run_blob_conversion_easy_nmt(session, local_storage_path, limit_docs: str = 50):
    """
    This function translates the data using easy nmt
    """
    model = EasyNMT("opus-mt")
    estimates_json = {"api_len": 0, "api_cost": 0}
    json_files_list = get_all_files_in_directory(local_storage_path, wildcard="*.json")
    # existing_docs_list = [doc.id for doc in session.query(DocumentInstances)]
    annotated_docs_list = set([doc.document_id for doc in session.query(BlobLvlAnnotations)])
    json_files_list = [
        json_file
        for json_file in json_files_list
        if os.path.splitext(os.path.basename(json_file))[0] in annotated_docs_list
    ]

    if limit_docs:
        json_files_list = json_files_list[:limit_docs]
    # json_translate_estimate
    for json_file in tqdm(json_files_list, desc="Processing jsons for translation task"):
        if not os.path.exists(f"{os.path.splitext(json_file)[0]}esy_nmt_de.json"):
            with open(json_file, "r") as f:
                doc_dict = json.load(f)

            for _, page_blobs in tqdm(
                doc_dict.items(), desc=f"Translating blobs for : {os.path.splitext(os.path.basename(json_file))[0]}"
            ):
                for blob in page_blobs:
                    blob["text"] = model.translate(blob["text"], target_lang="de")

            with open(f"{os.path.splitext(json_file)[0]}esy_nmt_de.json", "w") as f:
                json.dump(doc_dict, f, indent=2)
    return estimates_json


def run_annotated_blob_conversion(session):
    pass


def run_compliance_items_translation():
    pass


def dump_esrs_data_todb(session, raw_data_file: str, esrs_text_csv: str):
    all_reqs = defaultdict(dict)
    with open(raw_data_file, "r") as f:
        raw_data_dict = json.load(f)
    for esrs_cat, esrs_items in tqdm(raw_data_dict.items(), desc="Processing ESRS"):
        esrs_cat = esrs_cat.strip("ESRS").strip()
        for _, esrs_sub_items in esrs_items.items():
            objective_requirements = esrs_sub_items.pop("objective_requirements", None)
            application_requirements = esrs_sub_items.pop("application_requirements", None)
            if objective_requirements:
                for para_id, objective_requirement in objective_requirements.items():
                    all_reqs[esrs_cat][para_id] = objective_requirement["text"]
                    if "sub_req" in objective_requirement:
                        subreqs = objective_requirement.pop("sub_req", None)
                        if subreqs:
                            for sname, subreq in subreqs.items():
                                all_reqs[esrs_cat][f"{para_id} {sname}"] = subreq["text"]
                                if "sub_sub_req" in subreq:
                                    sub_sub_reqs = subreq.pop("sub_sub_req", None)
                                    if sub_sub_reqs:
                                        for ssname, sub_sub_req in sub_sub_reqs.items():
                                            ssname = f"({ssname.lower()})"
                                            all_reqs[esrs_cat][f"{para_id} {sname} {ssname}"] = sub_sub_req["text"]

            if application_requirements:
                for para_id, application_requirement in application_requirements.items():
                    para_id = para_id.replace(" ", "")
                    all_reqs[esrs_cat][para_id] = application_requirement["text"]
                    if "sub_req" in application_requirement:
                        subreqs = application_requirement.pop("sub_req", None)
                        if subreqs:
                            for sname, subreq in subreqs.items():
                                all_reqs[esrs_cat][f"{para_id} {sname}"] = subreq["text"]
                                if "sub_sub_req" in subreq:
                                    sub_sub_reqs = subreq.pop("sub_sub_req", None)
                                    if sub_sub_reqs:
                                        for ssname, sub_sub_req in sub_sub_reqs.items():
                                            ssname = f"({ssname.lower()})"
                                            all_reqs[esrs_cat][f"{para_id} {sname} {ssname}"] = sub_sub_req["text"]

    query = select(RptRequirementsMapping, StandardsList).join(
        RptRequirementsMapping, RptRequirementsMapping.standard == StandardsList.id
    )

    mapping_data = session.execute(query).all()

    for row in tqdm(mapping_data, desc="Processing all sources"):
        text = ""
        section_name = row[0].source
        std_mapping_id = row[0].standard
        esrs_family = row[1].family
        esrs_category_id = row[1].name.strip("ESRS").strip()
        all_req_keys = all_reqs[esrs_category_id].keys()
        if row[1].family == "esrs_v2":
            section_names = section_name.replace("+", ";").replace(",", ";").split(";")
            for section in section_names:
                section_clean = section.strip().strip(esrs_category_id).strip(".").strip()

                section_clean = "55 (c)" if section_clean == "55(c)" else section_clean

                matched_sec_name, _, _ = fuzzy_match_strings(section_clean, all_req_keys, threshold=97)

                text += all_reqs[esrs_category_id][matched_sec_name]
        elif row[1].family == "esrs":
            csv_data = pd.read_csv(esrs_text_csv)
            text_data = csv_data[(csv_data.source == row[0].source) & (csv_data.href.notnull())].relevant_text
            if isinstance(text_data.values[0], str):
                text = text_data.values[0]

        if text:
            esrs_req_map_model = EsrsReqMappingModel(
                esrs_category_id=esrs_category_id,
                esrs_family=esrs_family,
                section_name=section_name,
                std_mapping_id=std_mapping_id,
                text=text,
            )
            esrs_req_map_table = EsrsReqMapping(
                esrs_category_id=esrs_req_map_model.esrs_category_id,
                esrs_family=esrs_req_map_model.esrs_family,
                section_name=esrs_req_map_model.section_name,
                std_mapping_id=esrs_req_map_model.std_mapping_id,
                text=esrs_req_map_model.text,
            )
            session.add(esrs_req_map_table)
            session.commit()


def process_esrs_data(session, storage_path: str, model_path: str, pdf2img_out_path: str):
    esrs_standards = session.query(StandardsList).filter_by(family="esrs").all()
    for esrs_standard in tqdm(esrs_standards, desc="Downloading and parsing esrs drafts"):
        destination_folder = os.path.join(storage_path, esrs_standard.id)
        if not os.path.exists(destination_folder):
            os.makedirs(destination_folder)
        response = requests.get(esrs_standard.href)

        if response.status_code == 200:
            destination_path = os.path.join(destination_folder, f"{esrs_standard.id}.pdf")

            with open(destination_path, "wb") as file:
                file.write(response.content)

            get_page_dict(pdf_path=destination_path, model_path=model_path, img_path=pdf2img_out_path)


def curate_blob_lvl_annotations(session, dataset_path):
    """
    This is for curating the page numbers for document_ref table
    """
    # update document_ref
    blob_lvl_annotations = session.query(BlobLvlAnnotations).all()
    for blob_lvl_annotation in tqdm(blob_lvl_annotations, desc="curating revisions"):
        doc = session.query(DocumentInstances).filter_by(id=blob_lvl_annotation.document_id).all()
        fvwr = session.query(ValuesWithRevisions).filter_by(id=blob_lvl_annotation.revision_id).all()[0]
        doc_refs = process_annotation_idx(fvwr.document_ref)
        if doc_refs:
            doc_refs, _ = doc_refs
            assert len(doc) == 1, f"Duplicate doc found for {blob_lvl_annotation.document_id}"
            json_path = os.path.join(dataset_path, f"{doc[0].company_id}/{doc[0].year}/{doc[0].type}/{doc[0].id}.json")
            if os.path.exists(json_path):
                with open(json_path, "r") as f:
                    json_data = json.load(f)
                for doc_ref in doc_refs:
                    try:
                        if (
                            blob_lvl_annotation.blob_text
                            == json_data[str(doc_ref)][blob_lvl_annotation.blob_id]["text"]
                        ):
                            blob_lvl_annotation.document_ref = str(doc_ref)
                            session.commit()
                            break
                    except IndexError:
                        print(f"{blob_lvl_annotation.blob_id} is not present in page {doc_ref}")


def create_data_rag(session):
    doc_list = []
    doc_instances = session.query(DocumentInstances).all()
    print(doc_instances)


def json_to_docx(session, storage_location, limit: None):
    doc_instances = session.query(DocumentInstances).all()
    values_w_revision = session.query(ValuesWithRevisions).all()
    values_w_revision = [fvwr.document_id for fvwr in values_w_revision]
    if limit:
        doc_paths = [
            doc.wz_storage_location.rsplit(".", 1)[0] + ".json" for doc in doc_instances if doc.id in values_w_revision
        ][:limit]
    else:
        doc_paths = [
            doc.wz_storage_location.rsplit(".", 1)[0] + ".json"
            for doc in tqdm(doc_instances, desc="getting file paths")
            if doc.id in values_w_revision
        ]

    for json_file in tqdm(doc_paths, desc="Processing pdf parsed jsons"):
        json_data_path = os.path.splitext(os.path.basename(json_file))[0]
        tgt_path = os.path.join(storage_location, f"{json_data_path}.docx")

        if os.path.exists(json_file) and not os.path.exists(tgt_path):
            # Create a new DOCX document
            doc = Document()

            doc_id = os.path.splitext(os.path.basename(json_file))[0]

            # Load JSON data from file
            with open(json_file, "r") as f:
                data = json.load(f)

            # Add content to the DOCX document based on JSON data
            for page_no, blob_list in data.items():
                for blob_idx, blob_dict in enumerate(blob_list):
                    try:
                        text = blob_dict["text"]
                        if blob_dict["class_name"] != "table":
                            text = re.sub(r"-\s+", "", text)
                            text = re.sub(r"\s+", " ", text)
                            text = re.sub(r"\s$|^\s", "", text)
                        doc.add_paragraph(f"{doc_id}-{page_no}-{blob_idx}:::")
                        doc.add_paragraph(f"{text}")

                    except ValueError:
                        text = re.sub(
                            r'[\x07\x08\n\x0b\x0c\r\x0e\x0f=&gt;)",%&*0-9;<>AMP;OQRSTU$K\x8c\x8d\x8e]', "", text
                        )
                        text = re.sub(r"[^a-zA-Z0-9\s]", "", text)
                        text = re.sub(r"[^\x20-\x7E]", "", text)
                        doc.add_paragraph(f"{text}")

                    doc.add_paragraph("\n")

            # Save the DOCX document
            doc.save(tgt_path)


def docx_to_json(session, src_docx_path, tgt_json_path):
    blob_info_pattern1 = r"\d+-\d+:::"

    doc_instances = session.query(DocumentInstances).all()
    doc_paths = [doc.wz_storage_location.rsplit(".", 1)[0] + ".json" for doc in doc_instances]

    original_json_dict = {}
    for json_file in tqdm(doc_paths, desc="Processing pdf parsed jsons"):
        if os.path.exists(json_file):
            with open(json_file, "r") as f:
                data = json.load(f)
            original_json_dict[os.path.splitext(os.path.basename(json_file))[0]] = data

    # Load the DOCX document
    for docx_file in tqdm(os.listdir(src_docx_path), desc="Processing docx files"):
        docx_file = os.path.join(src_docx_path, docx_file)
        base_file_name = os.path.splitext(os.path.basename(docx_file))[0]
        doc = Document(docx_file)
        data = {}
        passage_idx_string = ""

        for paragraph in doc.paragraphs:
            if re.findall(blob_info_pattern1, paragraph.text):
                passage_idx_string = re.findall(blob_info_pattern1, paragraph.text)[0]
                passage_idx_string = passage_idx_string.strip(":::")
                page_blob_info = passage_idx_string.split("-")
                page_no, blob_idx = int(page_blob_info[0]), int(page_blob_info[1])
                continue
            # Split the paragraph text based on ':::'
            if passage_idx_string:
                parts = paragraph.text.split(":::")
                if len(parts) == 2:
                    original_json_dict[base_file_name[:-3]][str(page_no)][blob_idx] = paragraph.text

        json_file = os.path.join(tgt_json_path, f"{base_file_name}.json")
        with open(json_file, "w") as f:
            # json.dump(original_json_dict[base_file_name[:-3]], f, indent=2)
            json.dump(original_json_dict[base_file_name], f, indent=2)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--config",
        default="/cluster/home/repo/my_llm_experiments/esrs_data_collection/srn_data_collector/main.yaml",
        type=str,
        help="Path to config",
    )
    parser.add_argument(
        "--cleanup_selected_table",
        type=bool,
        help="Whether the clear blob annotations table",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    config = yaml.safe_load(open(args.config, "r"))

    raw_esrs_req_path = config.pop("raw_esrs_req_path")
    esrs_text_csv = config.pop("esrs_text_csv")
    json_to_docx_dest = config.pop("json_to_docx_dest")
    json_to_docx_docs_limit = config.pop("json_to_docx_docs_limit")
    dataset_base_path = config.pop("dataset_base_path")
    db_file_path = config.pop("db_file_path")
    docx_src = config.pop("docx_src")
    docx_to_json_dest = config.pop("docx_to_json_dest")
    gpt_reponse_file_path = config.pop("gpt_reponse_file_path")
    if not os.path.exists(gpt_reponse_file_path):
        os.mkdir(gpt_reponse_file_path)

    logger = setup_logger(log_file=config.pop("log_dir"))

    # clean up table
    if args.cleanup_selected_table:
        engine = create_engine(f"sqlite:///{db_file_path}")
        metadata = MetaData()
        metadata.reflect(bind=engine)
        table_name = config.pop("table_to_be_cleared")
        table = Table(table_name, metadata, autoload_with=engine)
        table.drop(engine)
        logger.info("Cleanup complete")

    # Create SQLite database
    engine = create_engine(f"sqlite:///{db_file_path}")

    # Build all the tables
    metadata = Base.metadata
    Base.metadata.create_all(bind=engine)

    # Creates session
    session = Session(engine)

    # Update companies data
    # collect_companies_data(session)

    # Update annotations data
    # collect_annotations(session)

    # collect documents metadata
    # collect_documents_metadata(session, local_storage_path=dataset_base_path)

    # collect compliance item data
    # collect_compliance_item_data(session)

    # collect reporting item data
    # collect_reporting_item_data(session)

    # collect blob level annotations from json files
    # error_stat_dict, cost_est_dict = collect_blob_lvl_annotations(
    #     session,
    #     logger,
    #     local_storage_path=dataset_base_path,
    #     error_stat_dict={},
    #     fuzzy_match_thresh=config.pop("fuzzy_match_thresh", None),
    #     estimate_cost=False,
    #     gpt_reponse_file_path=gpt_reponse_file_path,
    # )

    # Calculate gpt3.5 cost for reports translation task/llama2 data translation
    # convert_pdf_est_json = run_blob_conversion(
    #     session,
    #     config.pop("model_name"),
    #     estimate_cost=config.pop("estimate_conversion_cost"),
    #     local_storage_path=dataset_base_path,
    #     limit_docs=config.pop("docs_limit"),
    # )

    # Run blob conversion using easy NMT
    # run_blob_conversion_easy_nmt(session,
    #                              local_storage_path=dataset_base_path)

    json_to_docx(session, json_to_docx_dest, limit=None)

    # docx_to_json(session, docx_src, docx_to_json_dest)

    # create srn_datapoint table containing disclosure requirements detailed data
    # collect_standards_list(session)

    # reporting req mapping information
    # collect_std_requirements(session)

    # download and convert esrs_data esrs nov 2022 draft
    # process_esrs_data(session,
    #                   storage_path=config.pop('esrs_drafts_data'),
    #                   model_path=config.pop('model_path'),
    #                   pdf2img_out_path=config.pop('pdf2img_out_path'))

    # curate_database
    # curate_blob_lvl_annotations(session, dataset_base_path)

    # Dump esrs requirement to database
    # dump_esrs_data_todb(session, raw_esrs_req_path, esrs_text_csv)

    # Create data for RAG implementation
    # create_data_rag(session)

    session.close()

    # with open(config.pop("err_stats_json"), "w") as f:
    #     json.dump(error_stat_dict, f, indent=4)

    # with open(config.pop("api_est_json"), "w") as f:
    #     json.dump(cost_est_dict, f, indent=4)

    # with open(config.pop("convert_pdf_est_json"), "w") as f:
    #     json.dump(convert_pdf_est_json, f, indent=4)


if __name__ == "__main__":
    main()
